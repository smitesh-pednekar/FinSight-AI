"""Document text extraction service."""

import os
from typing import Optional
from pathlib import Path
import logging

from pypdf import PdfReader
from docx import Document as DocxDocument
from PIL import Image
import pytesseract

logger = logging.getLogger(__name__)


class DocumentExtractor:
    """Extract text from various document formats."""
    
    @staticmethod
    def extract_text(file_path: str, mime_type: Optional[str] = None) -> tuple[str, int]:
        """
        Extract text from document.
        
        Args:
            file_path: Path to document file
            mime_type: MIME type of file
            
        Returns:
            Tuple of (extracted_text, page_count)
        """
        file_path_obj = Path(file_path)
        
        if not file_path_obj.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = file_path_obj.suffix.lower()
        
        try:
            if extension == ".pdf" or (mime_type and "pdf" in mime_type):
                return DocumentExtractor._extract_from_pdf(file_path)
            elif extension in [".docx", ".doc"] or (mime_type and "word" in mime_type):
                return DocumentExtractor._extract_from_docx(file_path)
            elif extension in [".png", ".jpg", ".jpeg", ".tiff"] or (mime_type and "image" in mime_type):
                return DocumentExtractor._extract_from_image(file_path)
            else:
                logger.warning(f"Unsupported file type: {extension}")
                return "", 0
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise
    
    @staticmethod
    def _extract_from_pdf(file_path: str) -> tuple[str, int]:
        """Extract text from PDF."""
        text_parts = []
        page_count = 0
        
        try:
            with open(file_path, "rb") as file:
                pdf_reader = PdfReader(file)
                page_count = len(pdf_reader.pages)
                
                logger.info(f"PDF has {page_count} pages. Starting text extraction...")
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text_parts.append(f"\n--- Page {page_num} ---\n{page_text}")
                    else:
                        logger.warning(f"No selectable text found on page {page_num}")
            
            full_text = "\n".join(text_parts).strip()
            
            if not full_text:
                logger.warning(f"No selectable text extracted from the entire PDF: {file_path}")
                # Future: Add OCR fallback here if pdf2image is available
            
            return full_text, page_count
        except Exception as e:
            logger.error(f"Failed to read PDF {file_path}: {str(e)}")
            raise ValueError(f"Could not read PDF file: {str(e)}")
    
    @staticmethod
    def _extract_from_docx(file_path: str) -> tuple[str, int]:
        """Extract text from DOCX."""
        doc = DocxDocument(file_path)
        
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)
        
        full_text = "\n".join(text_parts)
        page_count = 1
        
        return full_text, page_count
    
    @staticmethod
    def _extract_from_image(file_path: str) -> tuple[str, int]:
        """Extract text from image using OCR."""
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            return text, 1
        except Exception as e:
            logger.error(f"OCR failed for {file_path}: {str(e)}")
            return "", 1
